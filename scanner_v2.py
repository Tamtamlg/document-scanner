import os
import win32com.client
import openpyxl
import xlrd
import socket
import webbrowser
from docx import Document
from odf.opendocument import load
from odf.text import P
from odf.table import Table, TableRow, TableCell


VERSION = '1.1.0'


def clean_text(text):
    return " ".join(text.split()).lower()


def search_in_tmp(file_path, search_texts):
    found_texts = []
    try:
        with open(file_path, "r", encoding="utf-8", errors="ignore") as f:
            text = clean_text(f.read())

            for phrase in search_texts:
                if clean_text(phrase) in text:
                    print(f"✅ Знайдено текст '{phrase}'")
                    found_texts.append(f"Знайдено текст '{phrase}'")
    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
    
    return found_texts


def search_in_xlsx(file_path, search_texts):
    found_texts = []
    try:
        wb = openpyxl.load_workbook(file_path, data_only=True)
        for sheet in wb.worksheets:
            for row in sheet.iter_rows():
                for cell in row:
                    if cell.value:
                        cell_text = clean_text(str(cell.value))
                        for phrase in search_texts:
                            if clean_text(phrase) in cell_text:
                                print(f"✅ Знайдено текст '{phrase}'")
                                found_texts.append(f"Знайдено текст '{phrase}' (лист: {sheet.title})")
    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
    return found_texts


def search_in_xls(file_path, search_texts):
    found_texts = []
    try:
        wb = xlrd.open_workbook(file_path)
        for sheet in wb.sheets():
            for row_idx in range(sheet.nrows):
                for col_idx in range(sheet.ncols):
                    cell_value = sheet.cell(row_idx, col_idx).value
                    if isinstance(cell_value, str):
                        cell_text = clean_text(cell_value)
                        for phrase in search_texts:
                            if clean_text(phrase) in cell_text:
                                print(f"✅ Знайдено текст '{phrase}'")
                                found_texts.append(f"Знайдено текст '{phrase}' (лист: {sheet.name})")
    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
    return found_texts


def search_in_doc(file_path, search_texts):
    found_texts = []
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(os.path.abspath(file_path))
        text = doc.Content.Text

        for phrase in search_texts:
            if phrase.lower() in text.lower():
                print(f"✅ Знайдено текст '{phrase}'")
                found_texts.append(f'Знайдено текст "{phrase}"')

        doc.Close(False)
        word.Quit()
    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
        try:
            word.Quit()
        except:
            pass
    return found_texts


def search_in_docx(file_path, search_texts):
    found_texts = []
    try:
        doc = Document(file_path)

        for para in doc.paragraphs:
            for phrase in search_texts:
                if phrase.lower() in para.text.lower():
                    print(f"✅ Знайдено текст '{phrase}'")
                    found_texts.append(f"Знайдено текст '{phrase}': {para.text}")

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    for phrase in search_texts:
                        if phrase.lower() in cell.text.lower():
                            print(f"✅ Знайдено текст '{phrase}'")
                            found_texts.append(f"Знайдено текст '{phrase}' в таблиці: {cell.text}")

    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
    return found_texts


def extract_text_from_odt(doc):
    TEXT_NODE = 3
    text = []
    for p in doc.getElementsByType(P):
        text.append("".join(node.data for node in p.childNodes if node.nodeType == TEXT_NODE))
    return " ".join(text)


def search_in_odt(file_path, search_texts):
    found_texts = []
    try:
        doc = load(file_path)
        text = clean_text(extract_text_from_odt(doc))

        for phrase in search_texts:
            if clean_text(phrase) in text:
                print(f"✅ Знайдено текст '{phrase}'")
                found_texts.append(f"Знайдено текст '{phrase}'")
    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")

    return found_texts


def extract_text_from_ods(doc):
    TEXT_NODE = 3
    text = []
    for table in doc.getElementsByType(Table):
        for row in table.getElementsByType(TableRow):
            for cell in row.getElementsByType(TableCell):
                paragraphs = cell.getElementsByType(P)
                for p in paragraphs:
                    cell_text = "".join(node.data for node in p.childNodes if node.nodeType == TEXT_NODE)
                    text.append(cell_text)
    return " ".join(text)


def search_in_ods(file_path, search_texts):
    found_texts = []
    try:
        doc = load(file_path)
        text = clean_text(extract_text_from_ods(doc))

        for phrase in search_texts:
            if clean_text(phrase) in text:
                print(f"✅ Знайдено текст '{phrase}'")
                found_texts.append(f"Знайдено текст '{phrase}'")
    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")

    return found_texts


def search_in_all_files(root_folder, search_texts):
    results = []

    for foldername, subfolders, filenames in os.walk(root_folder):
        for filename in filenames:
            file_path = os.path.join(foldername, filename)

            if filename.lower().endswith(".docx"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_docx(file_path, search_texts)
            elif filename.lower().endswith(".doc"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_doc(file_path, search_texts)
            elif filename.lower().endswith(".xlsx"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_xlsx(file_path, search_texts)
            elif filename.lower().endswith(".xls"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_xls(file_path, search_texts)
            elif filename.lower().endswith(".odt"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_odt(file_path, search_texts)
            elif filename.lower().endswith(".ods"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_ods(file_path, search_texts)
            elif filename.lower().endswith(".tmp"):
                print(f"🔍 Перевіряю: {file_path}")
                found_texts = search_in_tmp(file_path, search_texts)
            else:
                continue

            if found_texts:
                results.append((file_path, found_texts))

    return results


def get_computer_name():
    computer_name = socket.gethostname()
    return computer_name


def get_search_phrases():
    while True:
        print("🔍 Виберіть режим пошуку:")
        print("1 - Шукати тільки фразу \"Таємно\"")
        print("2 - Шукати фрази \"Для службового користування\", \"Таємно\"")
        print("3 - Ввести свої фрази вручну")

        choice = input("Ваш вибір (1, 2 або 3): ").strip()

        if choice == "1":
            return ["Таємно"]
        elif choice == "2":
            return ["Для службового користування", "Таємно"]
        elif choice == "3":
            while True:
                phrases_input = input("Введіть фрази через кому: ").strip()
                return [phrase.strip() for phrase in phrases_input.split(",") if phrase.strip()]
                if len(search_phrases) >= 1:
                    break
                else:
                    print("❌ Треба ввести хоча б одну фразу!")
            break
        else:
            print("❌ Невірний вибір! Спробуйте ще раз.\n")


if __name__ == "__main__":
    print(f"document_scanner v{VERSION} підтримує формати doc, docx, xls, xlsx, odt, ods\n")

    search_phrases = get_search_phrases()
    disks = input("Введіть букву диска для пошуку (наприклад, CDE): ").strip().upper()

    all_results = {}
    computer_name = get_computer_name()

    for disk in disks:
        root_folder = f"{disk}:\\"

        if os.path.exists(root_folder):
            print(f"\n🔍 Починаємо пошук на диску {root_folder} ...")
            results = search_in_all_files(root_folder, search_phrases)

            all_results[disk] = results

            html_file = f"{computer_name}_{disk}.html"
            with open(html_file, "w", encoding="utf-8") as f:
                f.write("""
                    <html>
                    <head>
                    <meta charset="utf-8">
                    """)
                f.write(f"<title>Диск {disk}</title>")
                f.write("""
                    <style>
                        body { font-family: Arial, sans-serif; margin: 20px; }
                        table { border-collapse: collapse; width: 100%; }
                        th, td { border: 1px solid #999; padding: 8px; text-align: left; }
                        th { background-color: #f2f2f2; cursor: pointer; }
                        tr:hover { background-color: #f9f9f9; }
                        .highlight { background-color: #ffff99; font-weight: bold; }
                        #searchInput { margin-bottom: 15px; padding: 5px; width: 300px; }
                    </style>
                    <script>
                        function sortTable(n) {
                            var table, rows, switching, i, x, y, shouldSwitch, dir, switchcount = 0;
                            table = document.getElementById("resultTable");
                            switching = true;
                            dir = "asc"; 
                            while (switching) {
                                switching = false;
                                rows = table.rows;
                                for (i = 1; i < (rows.length - 1); i++) {
                                    shouldSwitch = false;
                                    x = rows[i].getElementsByTagName("TD")[n];
                                    y = rows[i + 1].getElementsByTagName("TD")[n];
                                    if (dir == "asc") {
                                        if (x.innerText.toLowerCase() > y.innerText.toLowerCase()) {
                                            shouldSwitch = true;
                                            break;
                                        }
                                    } else if (dir == "desc") {
                                        if (x.innerText.toLowerCase() < y.innerText.toLowerCase()) {
                                            shouldSwitch = true;
                                            break;
                                        }
                                    }
                                }
                                if (shouldSwitch) {
                                    rows[i].parentNode.insertBefore(rows[i + 1], rows[i]);
                                    switching = true;
                                    switchcount ++;
                                } else {
                                    if (switchcount == 0 && dir == "asc") {
                                        dir = "desc";
                                        switching = true;
                                    }
                                }
                            }
                        }

                        function filterTable() {
                            var input, filter, table, tr, td, i, j, txtValue, found;
                            input = document.getElementById("searchInput");
                            filter = input.value.toLowerCase();
                            table = document.getElementById("resultTable");
                            tr = table.getElementsByTagName("tr");

                            for (i = 1; i < tr.length; i++) {
                                tr[i].style.display = "none";
                                td = tr[i].getElementsByTagName("td");
                                found = false;
                                for (j = 0; j < td.length; j++) {
                                    if (td[j] && td[j].innerText.toLowerCase().indexOf(filter) > -1) {
                                        found = true;
                                        break;
                                    }
                                }
                                if (found) {
                                    tr[i].style.display = "";
                                }
                            }
                        }
                    </script>
                    </head>
                    <body>
                    """)
                f.write(f"<h2>Результати перевірки {computer_name} (на  диску {disk} знайдено файлів: {len(results)})</h2>")
                f.write("""
                    <input type="text" id="searchInput" onkeyup="filterTable()" placeholder="Пошук по таблиці...">

                    <table id="resultTable">
                    <tr>
                        <th onclick="sortTable(0)">Файл</th>
                        <th onclick="sortTable(1)">Знайдені фрази</th>
                    </tr>
                    """)

                for file, texts in results:
                    formatted_texts = []
                    for text in texts:
                        for phrase in search_phrases:
                            highlighted = f"<span class='highlight'>{phrase}</span>"
                            text = text.replace(phrase, highlighted)
                        formatted_texts.append(text)
                    joined_texts = "<br>".join(formatted_texts)

                    f.write(f"<tr><td>{file}</td><td>{joined_texts}</td></tr>\n")

                f.write("""
            </table>
            </body>
            </html>
            """)

            print(f"✅ Знайдено файлів на диску {disk}: {len(results)}. Результати збережено в {html_file}")

            webbrowser.open(f"file://{os.path.abspath(html_file)}")

        else:
            print(f"❌ Помилка: диск {disk}:\\ не знайдено!")

    input("\n🔹 Натисніть Enter, для завершення роботи програми...")
