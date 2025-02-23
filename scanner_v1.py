import os
import win32com.client
from docx import Document


def search_in_doc(file_path, search_texts):
    try:
        word = win32com.client.Dispatch("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0

        doc = word.Documents.Open(os.path.abspath(file_path))
        text = doc.Content.Text.lower()

        found = any(phrase.lower() in text for phrase in search_texts)
        doc.Close(False)
        word.Quit()
        return found

    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
        try:
            word.Quit()
        except:
            pass
        return False

def search_in_docx(file_path, search_texts):
    try:
        doc = Document(file_path)
        search_texts = [phrase.lower() for phrase in search_texts]

        for para in doc.paragraphs:
            if any(phrase in para.text.lower() for phrase in search_texts):
                return True

        for table in doc.tables:
            for row in table.rows:
                for cell in row.cells:
                    if any(phrase in cell.text.lower() for phrase in search_texts):
                        return True

    except Exception as e:
        print(f"❌ Помилка обробки {file_path}: {e}")
    return False



def search_in_word_file(file_path, search_text):
    if file_path.lower().endswith(".docx"):
        return search_in_docx(file_path, search_text)
    elif file_path.lower().endswith(".doc"):
        return search_in_doc(file_path, search_text)
    return False


def search_in_all_docs(root_folder, search_text, output_file="results.txt"):
    found_files = []
    
    with open(output_file, "a", encoding="utf-8") as f:
        f.write(f"Результати пошуку '{search_text}' в документах:\n\n")

        for dirpath, _, filenames in os.walk(root_folder):
            for file in filenames:
                if file.lower().endswith((".doc", ".docx")):
                    file_path = os.path.join(dirpath, file)
                    print(f"🔍 Перевіряю: {file_path}")

                    if search_in_word_file(file_path, search_text):
                        print(f'✅ Знайдено в файлі: {file_path}')
                        f.write(f"{file_path}\n")
                        found_files.append(file_path)
        f.write("\n\n")
    
    print("✅ Пошук завершено.")
    print(f"📄 Результати записано в {output_file}")
    return found_files


if __name__ == "__main__":

    search_phrases = ["Для службового користування", "Таємно"]
    disks = input("Введіть букву диска для пошуку (наприклад, CDE): ").strip().upper()

    for disk in disks:
        root_folder = f"{disk}:\\"
        
        if os.path.exists(root_folder):
            print(f"\n🔍 Починаємо пошук на диску {root_folder} ...")
            search_in_all_docs(root_folder, search_phrases)
        else:
            print("❌ Помилка: диск не знайдено!")


    input("\n🔹 Натисніть Enter, для завершення роботи програми...")